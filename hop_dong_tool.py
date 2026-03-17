#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
=============================================================
  Happy Smart Light — Công cụ Tạo Hợp Đồng Mua Bán v2.0
  Input : Hóa đơn điện tử (XML / PDF)
  Output: Hợp đồng mua bán (.docx)
=============================================================
"""

import os, sys, json, re, threading
import xml.etree.ElementTree as ET
from datetime import date, datetime
from pathlib import Path

# Try importing utils if available (for centralized path handling)
try:
    from utils.paths import get_resource_path
except ImportError:
    # Fallback to simple logic if not running from main tool
    def get_resource_path(rel):
        import sys
        base = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
        return Path(base) / rel

# ── Đảm bảo PyQt6 có mặt TRƯỚC khi định nghĩa các class ──
try:
    import PyQt6
except ImportError:
    import subprocess
    print("📦  Đang cài PyQt6, vui lòng chờ…")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "PyQt6"])
    print("✅  Cài xong. Khởi động lại…")
    os.execv(sys.executable, [sys.executable] + sys.argv)

# ── Optional: python-docx ──────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as lxml_etree
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

# ── Seller constants ───────────────────────────────────────
SELLER = {
    "name":           "CÔNG TY TNHH THƯƠNG MẠI VÀ CÔNG NGHỆ HAPPY SMART LIGHT",
    "address":        "42 Hà Đức Trọng, Phường Bà Rịa, Thành phố Hồ Chí Minh",
    "tax_code":       "3502535621",
    "representative": "NGUYỄN DUY BẰNG",
    "title":          "Giám đốc",
    "accounts": [
        {
            "number": "7294949999",
            "holder": "CÔNG TY HAPPY SMART LIGHT",
            "bank":   "MB Bank (Ngân hàng Quân Đội)",
        },
        {
            "number": "72949488",
            "holder": "CÔNG TY HAPPY SMART LIGHT",
            "bank":   "Techcombank (Ngân hàng Kỹ Thương Việt Nam)",
        },
    ],
}

# ── Versioning ─────────────────────────────────────────────
def load_version():
    try:
        conf_path = get_resource_path("config.json")
        if conf_path.exists():
            with open(conf_path, "r", encoding="utf-8") as f:
                return json.load(f).get("version", "2.1.0")
    except:
        pass
    return "2.1.0"

VERSION = load_version()

# ── Utilities ──────────────────────────────────────────────
def fmt(n) -> str:
    try:
        return f"{int(n):,}".replace(",", ".")
    except:
        return str(n)

def parse_int(s: str) -> int:
    return int(re.sub(r"[^\d]", "", str(s)) or "0")

def split_half(total: int):
    a = total // 2
    return a, total - a

def parse_date_str(s: str) -> date:
    for fmt_str in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
        try:
            return datetime.strptime(s.strip(), fmt_str).date()
        except:
            pass
    return date.today()

# ── Invoice parsing ────────────────────────────────────────
class Invoice:
    def __init__(self):
        self.no = ""; self.serial = ""; self.inv_date = ""
        self.payment_method = ""
        self.buyer_name = ""; self.buyer_tax = ""; self.buyer_address = ""
        self.items = []
        self.total_before_tax = 0
        self.total_tax = 0
        self.total_payment = 0
        self.total_words = ""

def parse_xml(path: str) -> Invoice:
    inv = Invoice()
    root = ET.parse(path).getroot()
    tc = root.find(".//TTChung")
    if tc is not None:
        inv.no             = tc.findtext("SHDon", "")
        inv.serial         = tc.findtext("KHHDon", "")
        inv.inv_date       = tc.findtext("NLap", "")
        inv.payment_method = tc.findtext("HTTToan", "")
    nm = root.find(".//NMua")
    if nm is not None:
        inv.buyer_name    = nm.findtext("Ten", "")
        inv.buyer_tax     = nm.findtext("MST", "")
        inv.buyer_address = nm.findtext("DChi", "")
    for h in root.findall(".//HHDVu"):
        before_tax = parse_int(h.findtext("ThTien", "0"))
        tax_rate_s = h.findtext("TSuat", "8%")
        tax_pct    = parse_int(tax_rate_s) / 100
        tax_amount = round(before_tax * tax_pct)
        inv.items.append({
            "stt":        h.findtext("STT", ""),
            "name":       h.findtext("THHDVu", ""),
            "unit":       h.findtext("DVTinh", ""),
            "qty":        h.findtext("SLuong", ""),
            "unit_price": parse_int(h.findtext("DGia", "0")),
            "before_tax": before_tax,
            "tax_rate":   tax_rate_s,
            "tax_amount": tax_amount,
            "total":      before_tax + tax_amount,
        })
    tt = root.find(".//TToan")
    if tt is not None:
        inv.total_before_tax = parse_int(tt.findtext("TgTCThue", "0"))
        inv.total_tax        = parse_int(tt.findtext("TgTThue",  "0"))
        inv.total_payment    = parse_int(tt.findtext("TgTTTBSo", "0"))
        inv.total_words      = tt.findtext("TgTTTBChu", "")
    return inv

def parse_pdf(path: str) -> Invoice:
    if not HAS_PDF:
        raise RuntimeError("Cần cài pdfplumber:\n  pip install pdfplumber")
    inv = Invoice()
    with pdfplumber.open(path) as pdf:
        text = "\n".join(p.extract_text() or "" for p in pdf.pages)
    def find(pattern, flags=0):
        m = re.search(pattern, text, re.IGNORECASE | flags)
        return m.group(1).strip() if m else ""
    inv.buyer_name    = find(r"Tên đơn vị[^:]*:\s*(.+)")
    inv.buyer_tax     = find(r"Mã số thuế[^:]*:\s*(\d[\d\-]*)")
    inv.buyer_address = find(r"Địa chỉ[^:]*:\s*(.+)")
    inv.total_words   = find(r"Số tiền viết bằng chữ[^:]*:\s*(.+)")
    inv.inv_date      = find(r"Ngày\s+\(date\)\s+(\d{2}/\d{2}/\d{4})")
    return inv

def parse_html(path: str) -> Invoice:
    if not HAS_BS4:
        raise RuntimeError("Cần cài beautifulsoup4 và lxml:\n  pip install beautifulsoup4 lxml")
    
    inv = Invoice()
    with open(path, "r", encoding="utf-8") as f:
        html_content = f.read()
    
    # Pre-clean non-breaking spaces and fragmentation
    html_content = html_content.replace("&nbsp;", " ")
    soup = BeautifulSoup(html_content, "lxml")
    
    # Extract flat text for regex matches
    text = soup.get_text(separator=" ").strip()
    flat_text = re.sub(r"\s+", " ", text)
    
    def find_val(pattern, src=flat_text):
        m = re.search(pattern, src, re.IGNORECASE)
        return m.group(1).strip() if m else ""

    # Basic info
    inv.no = find_val(r"Số \(No\)\s*:\s*(\d+)")
    inv.serial = find_val(r"Ký hiệu \(Serial\)\s*:\s*([A-Z0-9]+)")
    
    # Date
    day = find_val(r"Ngày \(date\)\s+(\d{1,2})")
    month = find_val(r"tháng \(month\)\s+(\d{1,2})")
    year = find_val(r"năm \(year\)\s+(\d{4})")
    if day and month and year:
        inv.inv_date = f"{year.zfill(4)}-{month.zfill(2)}-{day.zfill(2)}"
    
    # Buyer info isolation
    buyer_m = re.search(r"Họ tên người mua hàng(.+?)Tên hàng hóa", flat_text, re.IGNORECASE)
    buyer_section = buyer_m.group(1) if buyer_m else flat_text
    inv.buyer_name = find_val(r"Tên đơn vị \(Company\)\s*:\s*(.+?)(?=\sMã số thuế|Địa chỉ|$)", buyer_section)
    inv.buyer_tax = find_val(r"Mã số thuế \(Tax code\)\s*:\s*(\d+)", buyer_section)
    
    # Robust address stopping barriers
    addr_stop = r"\sĐiện thoại|\sSố tài khoản|\sHình thức thanh toán|\sMã số thuế|\sFax|\sEmail|$"
    inv.buyer_address = find_val(fr"Địa chỉ \(Address\)\s*:\s*(.+?)(?={addr_stop})", buyer_section)
    
    inv.payment_method = find_val(r"Hình thức thanh toán \(Payment method\)\s*:\s*(.+?)(?=\sSố tài khoản|$)")
    inv.total_words = find_val(r"Số tiền viết bằng chữ \(Amount in words\)\s*:\s*(.+?)(?=\.)")
    
    # Items extraction - Filter out column numbering
    rows = soup.find_all("tr")
    found_table = False
    for row in rows:
        tds = row.find_all("td")
        cols = [td.get_text(strip=True) for td in tds]
        if not cols: continue
        if any("STT" in c for c in cols) and any("hàng hóa" in c.lower() for c in cols):
            found_table = True
            continue
        if found_table:
            cell_vals = [c for c in cols if c]
            if not cell_vals: continue
            if cell_vals[0].isdigit() and len(cell_vals[0]) < 4:
                if len(cell_vals) >= 7 and cell_vals[1] == "2" and cell_vals[2] == "3": continue
                if len(cell_vals) >= 4:
                    item = {
                        "stt":        cell_vals[0],
                        "name":       cell_vals[1],
                        "unit":       cell_vals[2] if len(cell_vals) > 2 else "",
                        "qty":        cell_vals[3] if len(cell_vals) > 3 else "",
                        "unit_price": parse_int(cell_vals[4]) if len(cell_vals) > 4 else 0,
                        "before_tax": parse_int(cell_vals[5]) if len(cell_vals) > 5 else 0,
                        "tax_rate":   cell_vals[6] if len(cell_vals) > 6 else "",
                        "tax_amount": parse_int(cell_vals[7]) if len(cell_vals) > 8 else 0,
                        "total":      parse_int(cell_vals[-1]),
                    }
                    inv.items.append(item)
            if any("Cộng" in c for c in cell_vals) or any("Total" in c for c in cell_vals): break

    # Totals - targeted extraction
    for row in rows:
        r_text = row.get_text(separator=" ", strip=True)
        if "(Total):" in r_text:
            tds = row.find_all("td")
            nums = []
            for td in tds:
                t = td.get_text(strip=True)
                if re.search(r"[\d\.,]{5,}", t):
                    nums.append(parse_int(t))
            if len(nums) >= 3:
                inv.total_before_tax = nums[0]
                inv.total_tax = nums[1]
                inv.total_payment = nums[2]
                break

    # Fallback absolute regex
    if not inv.total_before_tax: inv.total_before_tax = parse_int(find_val(r"\(Total\):.*?([\d\.,]{7,})"))
    if not inv.total_tax: inv.total_tax = parse_int(find_val(r"\(VAT Amount\).*?([\d\.,]{5,})"))
    if not inv.total_payment:
        tp_m = re.search(r"Total of payment\):.*?([\d\.,]{7,})", flat_text, re.IGNORECASE)
        if tp_m:
            inv.total_payment = parse_int(tp_m.group(1))
        else:
             tp_m = re.search(r"([\d\.,]{7,})\s*Số tiền viết bằng chữ", flat_text, re.IGNORECASE)
             if tp_m: inv.total_payment = parse_int(tp_m.group(1))

    return inv

# ── Tax code API ───────────────────────────────────────────
def lookup_mst(mst: str) -> dict:
    if not HAS_REQUESTS:
        return {"error": "Cần cài requests:\n  pip install requests"}
    try:
        url = f"https://api.xinvoice.vn/gdt-api/tax-payer/{mst.strip()}"
        headers = {"Accept": "application/json"}
        r = requests.get(url, headers=headers, timeout=8)
        if r.status_code == 404:
            return {"error": "Không tìm thấy MST này"}
        if r.status_code == 200:
            j = r.json()
            raw_status = j.get("status", "")
            is_active  = "đang hoạt động" in raw_status.lower()
            return {
                "tax_id":      j.get("taxID", mst.strip()),  # Trả về mã số thuế từ API hoặc giá trị nhập
                "name":        j.get("name", ""),
                "address":     j.get("address", ""),
                "status":      "active" if is_active else "inactive",
                "status_text": raw_status,
            }
        return {"error": f"HTTP {r.status_code}"}
    except Exception as e:
        return {"error": str(e)}

# ── DOCX generation ────────────────────────────────────────
def set_cell(cell, text, bold=False, italic=False, sz=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, bg=None):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run(text)
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(sz)
    run.font.name  = "Times New Roman"
    if bg:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg)
        tcPr.append(shd)

def generate_docx(data: dict, out_path: str):
    if not HAS_DOCX:
        raise RuntimeError("Cần cài python-docx:\n  pip install python-docx")

    doc = Document()
    sec = doc.sections[0]
    sec.page_height   = Cm(29.7); sec.page_width   = Cm(21)
    sec.left_margin   = Cm(2.5);  sec.right_margin = Cm(2)
    sec.top_margin    = Cm(2);    sec.bottom_margin = Cm(2)

    # ── Watermark logo ───────────────────────────────────────
    logo_path = get_resource_path("logo.png")
    if logo_path.exists():
        header = sec.header
        header.is_linked_to_previous = False
        hdr_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hdr_para.clear()
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = hdr_para.add_run()
        run.add_picture(str(logo_path), width=Cm(8))

        WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        A  = "http://schemas.openxmlformats.org/drawingml/2006/main"
        inline_elem = hdr_para._p.find(".//{%s}inline" % WP)
        if inline_elem is not None:
            cx_val = inline_elem.find("{%s}extent" % WP).get("cx")
            cy_val = inline_elem.find("{%s}extent" % WP).get("cy")
            pos_x  = str((7560000 - int(cx_val)) // 2)
            pos_y  = str((10692000 - int(cy_val)) // 2)

            graphic_elem = inline_elem.find(".//{%s}graphic" % A)
            graphic_xml  = lxml_etree.tostring(graphic_elem, encoding="unicode") if graphic_elem is not None else ""

            anchor_xml = (
                f'<wp:anchor distT="0" distB="0" distL="0" distR="0" '
                f'simplePos="0" relativeHeight="251658240" behindDoc="1" '
                f'locked="0" layoutInCell="1" allowOverlap="1" '
                f'xmlns:wp="{WP}">'
                f'<wp:simplePos x="0" y="0"/>'
                f'<wp:positionH relativeFrom="page"><wp:posOffset>{pos_x}</wp:posOffset></wp:positionH>'
                f'<wp:positionV relativeFrom="page"><wp:posOffset>{pos_y}</wp:posOffset></wp:positionV>'
                f'<wp:extent cx="{cx_val}" cy="{cy_val}"/>'
                f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
                f'<wp:wrapNone/>'
                f'<wp:docPr id="100" name="WatermarkLogo"/>'
                f'<wp:cNvGraphicFramePr/>'
                f'{graphic_xml}'
                f'</wp:anchor>'
            )
            anchor_elem = lxml_etree.fromstring(anchor_xml)
            drawing = inline_elem.getparent()
            drawing.remove(inline_elem)
            drawing.append(anchor_elem)

        a_blip = hdr_para._p.find(".//{%s}blip" % A)
        if a_blip is not None:
            alpha_mod = lxml_etree.SubElement(a_blip, "{%s}alphaModFix" % A)
            alpha_mod.set("amt", "25000")

    FONT = "Times New Roman"
    sign_year = data["sign_date"].year

    def add_para(text="", bold=False, italic=False, sz=13,
                 align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after  = Pt(after)
        r = para.add_run(text)
        r.bold      = bold
        r.italic    = italic
        r.font.size = Pt(sz)
        r.font.name = FONT
        return para

    # ── Header ──────────────────────────────────────────────
    add_para("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
             bold=True, sz=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para("Độc lập – Tự do – Hạnh phúc",
             bold=True, sz=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para("---o0o---", sz=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para("HỢP ĐỒNG MUA BÁN HÀNG HÓA",
             bold=True, sz=15, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(f"Số: {data['contract_no']}/{sign_year}/HDMB",
             bold=True, sz=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    for line in [
        "- Căn cứ Bộ luật Dân sự số 91/2015/QH13 ngày 24/11/2015 và các văn bản pháp luật liên quan;",
        "- Căn cứ Luật Thương mại số 36/2005/QH11 ngày 14/06/2005 và các văn bản pháp luật liên quan;",
        "- Căn cứ vào nhu cầu và khả năng của các bên;",
    ]:
        add_para(line, sz=13, after=2)

    sd = data["sign_date"]
    add_para(f"\nHôm nay, ngày {sd.day} tháng {sd.month} năm {sd.year}, chúng tôi gồm có:", after=8)

    # ── Bên A ────────────────────────────────────────────────
    s = SELLER
    add_para(f'BÊN BÁN ("BÊN A"): {s["name"]}.', bold=True, after=2)
    add_para(f'Địa chỉ\t\t: {s["address"]}', after=2)
    add_para(f'Mã số thuế\t: {s["tax_code"]}', after=2)
    add_para(f'Đại diện\t: Ông {s["representative"]}'
             f'\t\t\tChức vụ : {s["title"]}.', after=10)

    # ── Bên B ────────────────────────────────────────────────
    b = data["buyer"]
    add_para(f'BÊN MUA ("BÊN B"): {b["name"]}', bold=True, after=2)
    add_para(f'Địa chỉ\t\t: {b["address"]}', after=2)
    add_para(f'Mã số thuế \t: {b["tax_code"]}', after=2)
    add_para(f'Đại diện \t: Ông {b["representative"]}'
             f'\t\t\tChức vụ  : {b["title"]}.', after=10)

    add_para("Trên cơ sở thỏa thuận, hai bên thống nhất ký kết hợp đồng mua bán hàng hóa "
             "với các điều khoản sau đây:", after=10)

    # ── Điều 1 ───────────────────────────────────────────────
    add_para("ĐIỀU 1. TÊN HÀNG - SỐ LƯỢNG - GIÁ TRỊ HỢP ĐỒNG",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para("Đơn vị tính: Việt Nam đồng",
             italic=True, sz=11, align=WD_ALIGN_PARAGRAPH.RIGHT, after=4)

    COL_W_DXA = [504, 3276, 576, 432, 1080, 1152, 576, 1080, 1152]
    HEADERS = [
        "STT", "Tên hàng hóa, dịch vụ", "ĐVT", "SL",
        "Đơn giá", "T.Tiền trước thuế", "TS%", "Tiền thuế", "Trị giá TT"
    ]
    tbl = doc.add_table(rows=1, cols=9)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = tbl._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tbl_pr)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"),    str(sum(COL_W_DXA)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    def set_col_width(cell, dxa):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"),    str(dxa))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    hdr_row = tbl.rows[0]
    for i, (h, w) in enumerate(zip(HEADERS, COL_W_DXA)):
        set_col_width(hdr_row.cells[i], w)
        set_cell(hdr_row.cells[i], h, bold=True, bg="D9E1F2")

    for item in data.get("items", []):
        row = tbl.add_row()
        vals = [
            str(item["stt"]), item["name"], item["unit"], str(item["qty"]),
            fmt(item["unit_price"]), fmt(item["before_tax"]),
            item["tax_rate"], fmt(item["tax_amount"]), fmt(item["total"]),
        ]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,  WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ]
        for i, (v, w, a) in enumerate(zip(vals, COL_W_DXA, aligns)):
            set_col_width(row.cells[i], w)
            set_cell(row.cells[i], v, align=a)

    for _ in range(3):
        row = tbl.add_row()
        for i, w in enumerate(COL_W_DXA):
            set_col_width(row.cells[i], w)
            set_cell(row.cells[i], "")

    tp  = data["total_payment"]
    tbt = data["total_before_tax"]
    ttv = data["total_tax"]
    sr  = tbl.add_row()
    for i, w in enumerate(COL_W_DXA):
        set_col_width(sr.cells[i], w)
    sr.cells[0].merge(sr.cells[3])
    set_cell(sr.cells[0], "Tổng cộng (Total):", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(sr.cells[4], "")
    set_cell(sr.cells[5], fmt(tbt), align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(sr.cells[6], "")
    set_cell(sr.cells[7], fmt(ttv), align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell(sr.cells[8], fmt(tp),  align=WD_ALIGN_PARAGRAPH.RIGHT)

    wr = tbl.add_row()
    for i, w in enumerate(COL_W_DXA):
        set_col_width(wr.cells[i], w)
    wr.cells[0].merge(wr.cells[8])
    set_cell(wr.cells[0],
             f"Số tiền viết bằng chữ: {data['total_words']}",
             italic=True, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()

    # ── Điều 2 ───────────────────────────────────────────────
    add_para("ĐIỀU 2. THANH TOÁN.", bold=True, after=4)
    pay_mode = data.get("pay_mode", "50_50")
    p1d  = data["pay1_date"]
    bank = data["bank"]

    if pay_mode == "100":
        add_para("Bên Mua phải thanh toán toàn bộ số tiền ghi tại Điều 1 ngay sau khi ký "
                 f"hợp đồng vào ngày {p1d.day}/{p1d.month}/{p1d.year}.", after=6)
    else:
        pct1  = 70 if pay_mode == "70_30" else 50
        pct2  = 100 - pct1
        inst1 = round(tp * pct1 / 100)
        inst2 = tp - inst1
        p2d   = data["pay2_date"]
        add_para("Bên Mua phải thanh toán cho bên Bán số tiền ghi tại Điều 1 của Hợp đồng 2 đợt:", after=2)
        add_para(f"Đợt 1: Thanh toán {pct1}% giá trị đơn hàng ({fmt(inst1)} đồng) "
                 f"ngay sau khi ký hợp đồng vào ngày {p1d.day}/{p1d.month}/{p1d.year}.", after=2)
        add_para(f"Đợt 2: Thanh toán {pct2}% số tiền còn lại ({fmt(inst2)} đồng) "
                 f"ngay khi nhận được hàng vào ngày {p2d.day}/{p2d.month}/{p2d.year}.", after=6)
    add_para("Bên Mua thanh toán cho Bên Bán theo hình thức chuyển khoản vào tài khoản "
             "của Bên Bán chi tiết như sau:", after=2)
    add_para(f"Số tài khoản: {bank['number']}", after=2)
    add_para(f"Chủ tài khoản: {bank['holder']}.", after=2)
    add_para(f"Mở tại: {bank['bank']}", after=10)

    # ── Điều 3 ───────────────────────────────────────────────
    dd = data["delivery_date"]
    add_para("ĐIỀU 3: THỜI GIAN, ĐỊA ĐIỂM VÀ PHƯƠNG THỨC GIAO HÀNG",
             bold=True, after=4)
    add_para(f"Việc bàn giao hàng hóa cho Bên Mua được thực hiện chậm nhất vào "
             f"ngày {dd.day}/{dd.month}/{dd.year}.", after=2)
    add_para(f"Địa điểm giao nhận hàng tại Trụ sở của Bên Mua: {b['address']}", after=2)
    add_para("Phương thức giao hàng: bên Bán sẽ giao hàng cho bên Mua qua ứng dụng giao hàng. "
             "Phí vận chuyển sẽ do bên Mua thanh toán.", after=10)

    # ── Điều 4 ───────────────────────────────────────────────
    add_para("ĐIỀU 4: TRÁCH NHIỆM CỦA CÁC BÊN", bold=True, after=4)
    for line in [
        "Bên Bán có nghĩa vụ giao hàng đúng thời gian, địa điểm, chất lượng và chủng loại "
        "theo quy định trong hợp đồng này, trừ trường hợp bất khả kháng.",
        "Bên Bán có nghĩa vụ cung cấp đầy đủ hóa đơn, chứng từ, tài liệu hợp lệ cho Bên Mua "
        "và chịu trách nhiệm pháp lý về tính hợp lệ của các hóa đơn, chứng từ, tài liệu đã giao.",
        "Bên Mua có trách nhiệm thanh toán và nhận hàng theo đúng thời gian đã quy định.",
        "Khi nhận hàng, bên mua có trách nhiệm kiểm nhận số lượng chủng loại hàng hóa tại chỗ "
        "và có video ghi nhận quá trình khui hàng. Nếu phát hiện hàng thiếu hoặc không đúng tiêu "
        "chuẩn chất lượng thì báo cho bên Bán ngay lập tức, yêu cầu bên Bán xác nhận. Hàng thiếu/"
        "không đúng loại mà bên Mua không có video khui hàng thì bên Bán không chịu trách nhiệm.",
        f"Trong trường hợp bên Bán giao hàng chậm trễ so với ngày dự kiến "
        f"({dd.day}/{dd.month}/{dd.year}) thì bên Bán có trách nhiệm thông báo cho bên Mua "
        f"để hai bên cùng tìm phương án giải quyết.",
    ]:
        add_para(line, after=3)
    doc.add_paragraph()

    # ── Điều 5 ───────────────────────────────────────────────
    add_para("ĐIỀU 5: BẤT KHẢ KHÁNG VÀ GIẢI QUYẾT TRANH CHẤP", bold=True, after=4)
    for line in [
        "Bất khả kháng có nghĩa là các sự kiện xảy ra một cách khách quan, không thể lường trước "
        "được và không thể khắc phục được mặc dù đã áp dụng mọi biện pháp cần thiết trong khả năng "
        "cho phép, một trong các Bên vẫn không có khả năng thực hiện được nghĩa vụ của mình theo "
        "Hợp đồng này, gồm nhưng không giới hạn ở: thiên tai, hỏa hoạn, lũ lụt, chiến tranh, can "
        "thiệp của chính quyền bằng vũ trang, cản trở giao thông vận tải và các sự kiện khác tương tự.",
        "Khi xảy ra sự kiện bất khả kháng, bên gặp phải bất khả kháng phải không chậm trễ, thông báo "
        "cho bên kia tình trạng thực tế, đề xuất phương án xử lý và nỗ lực giảm thiểu tổn thất, "
        "thiệt hại đến mức thấp nhất có thể.",
        "Trong quá trình thực hiện hợp đồng, nếu có vướng mắc từ bất kỳ bên nào, hai bên sẽ cùng nhau "
        "giải quyết trên tinh thần hợp tác. Trong trường hợp không tự giải quyết được, hai bên thống "
        "nhất đưa ra giải quyết tại Tòa án có thẩm quyền. Phán quyết của toà án là quyết định cuối cùng, "
        "có giá trị ràng buộc các bên. Bên thua phải chịu toàn bộ các chi phí giải quyết tranh chấp.",
    ]:
        add_para(line, after=3)
    doc.add_paragraph()

    # ── Điều 6 ───────────────────────────────────────────────
    add_para("ĐIỀU 6: ĐIỀU KHOẢN CHUNG", bold=True, after=4)
    for line in [
        "Việc thay đổi tên của Bên Mua trên hợp đồng và/hoặc trên hóa đơn sẽ không được chấp nhận.",
        "Hợp đồng này có giá trị thay thế mọi giao dịch, thỏa thuận trước đây của hai bên. "
        "Mọi sự bổ sung, sửa đổi hợp đồng này đều phải có sự đồng ý bằng văn bản của hai bên.",
        "Trừ các trường hợp được quy định ở trên, Hợp đồng này không thể bị hủy bỏ nếu không có "
        "thỏa thuận bằng văn bản của các bên.",
        "Hợp đồng này được lập thành 02 (hai) bản. Mỗi bên giữ 01 (một) bản, có giá trị pháp lý như nhau.",
        "Hợp đồng này có hiệu lực kể từ ngày ký và được coi là đã thanh lý khi Bên B đã nhận đủ tiền "
        "và Bên A đã nhận hàng.",
    ]:
        add_para(line, after=3)

    # ── Ký tên ───────────────────────────────────────────────
    doc.add_paragraph(); doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_info = [
        {"header": "Đại diện bên Mua", "sub": "(Ký và đóng dấu)", "name": ""},
        {"header": "Đại diện bên Bán", "sub": "(Ký và đóng dấu)", "name": "Nguyễn Duy Bằng"},
    ]
    for col_idx, info in enumerate(sig_info):
        c = sig.rows[0].cells[col_idx]
        c.text = ""
        para = c.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(info["header"])
        r.bold = True; r.font.size = Pt(13); r.font.name = FONT
        r.add_break()
        r2 = para.add_run(info["sub"])
        r2.font.size = Pt(13); r2.font.name = FONT
        for _ in range(5):
            c.add_paragraph()
        if info["name"]:
            np = c.add_paragraph()
            np.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rn = np.add_run(info["name"])
            rn.bold = True; rn.font.size = Pt(13); rn.font.name = FONT

    doc.save(out_path)


# ══════════════════════════════════════════════════════════
#  PyQt6 UI  —  Dark theme · Neon Magenta #e020d0 + Cyan #00c8f0
# ══════════════════════════════════════════════════════════
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QScrollArea, QFrame,
    QFileDialog, QMessageBox, QTableWidget, QTableWidgetItem,
    QRadioButton, QButtonGroup, QHeaderView, QComboBox,
    QDialog, QFormLayout, QAbstractItemView, QCalendarWidget, QDateEdit, QListView,
)
from PyQt6.QtCore import Qt, QDate, QThread, pyqtSignal, QSize
from PyQt6.QtGui  import QFont, QPixmap, QIcon, QColor

# ── Palette ────────────────────────────────────────────────
_PINK   = "#e020d0"
_CYAN   = "#00c8f0"
_BG     = "#0a0a14"
_CARD   = "#111120"
_BORDER = "#1e1e38"
_TEXT   = "#e8e8ff"
_DIM    = "#6868a0"
_ACCENT = "#16162a"
_GREEN  = "#00e87a"
_WARN   = "#ffaa00"

_QSS = f"""
/* ── Global ─────────────────────────────────────────── */
QMainWindow, QWidget {{ background:{_BG}; color:{_TEXT}; font-family:'Segoe UI',sans-serif; }}
QDialog  {{ background:{_CARD}; }}

/* ── Card frame ─────────────────────────────────────── */
QFrame#card {{
    background:{_CARD};
    border:1px solid {_BORDER};
    border-radius:10px;
}}
QLabel#card_hdr {{
    background:qlineargradient(x1:0,y1:0,x2:1,y2:0,
        stop:0 {_CYAN}28, stop:1 {_PINK}20);
    color:{_CYAN};
    border-top-left-radius:9px;
    border-top-right-radius:9px;
    border-bottom:1px solid {_BORDER};
    padding:9px 14px;
    font-weight:bold; font-size:13px;
}}

/* ── Labels ─────────────────────────────────────────── */
QLabel {{ color:{_TEXT}; font-size:12px; background:transparent; }}
QLabel#dim {{ color:{_DIM}; font-size:11px; }}
QLabel#ok  {{ color:{_GREEN}; font-size:11px; font-weight:bold; }}
QLabel#bad {{ color:{_PINK};  font-size:11px; font-weight:bold; }}
QLabel#info{{ color:{_CYAN};  font-size:11px; font-weight:bold; }}
QLabel#warn{{ color:{_WARN};  font-size:11px; font-weight:bold; }}
QLabel#total_amount {{ color:{_GREEN}; font-size:14px; font-weight:bold; }}
QLabel#total_words  {{ color:#00bb60;  font-size:11px; font-style:italic; }}
QLabel#suffix {{ color:{_DIM}; font-size:12px; }}

/* ── Inputs ─────────────────────────────────────────── */
QLineEdit {{
    background:{_ACCENT}; border:1px solid {_BORDER};
    border-radius:6px; color:{_TEXT};
    padding:6px 10px; font-size:12px;
    selection-background-color:{_CYAN}55;
}}
QLineEdit:focus  {{ border:1px solid {_CYAN}; }}
QLineEdit:hover  {{ border:1px solid {_CYAN}88; }}
QLineEdit:disabled {{ background:#0e0e1e; color:{_DIM}; }}

/* ── ComboBox ───────────────────────────────────────── */
QComboBox {{
    background:{_ACCENT}; border:1px solid {_CYAN}55;
    border-radius:6px; color:{_TEXT};
    padding:6px 10px; font-size:12px;
    min-width:80px;
    font-family: 'Segoe UI', sans-serif;
}}
QComboBox:focus  {{ border:1px solid {_CYAN}; }}
QComboBox:hover  {{ border:1px solid {_CYAN}; }}
QComboBox::drop-down {{
    border:none; background:transparent; width:28px;
    subcontrol-origin:border; subcontrol-position:right center;
}}
QComboBox::down-arrow {{
    image:none; border-left:5px solid transparent;
    border-right:5px solid transparent;
    border-top:6px solid {_CYAN}; width:0; height:0;
    margin-right:8px;
}}
QComboBox QAbstractItemView {{
    background-color: #000000;
    border: 1px solid {_CYAN};
    color:{_TEXT}; selection-background-color:{_CYAN}33;
    outline:none;
    font-size:12px;
}}
QComboBox QAbstractItemView::item {{
    font-size:12px;
    padding:6px 10px;
    font-family: 'Segoe UI', sans-serif;
    color: {_TEXT};
    background-color: #000000;
    min-height: 24px;
}}
QComboBox QAbstractItemView::item:selected {{
    background-color: {_CYAN}55;
    color: {_TEXT};
    font-weight: bold;
}}

/* ── Date Edit ───────────────────────────────────────── */
QDateEdit {{
    background:{_ACCENT}; border:1px solid {_CYAN}55;
    border-radius:6px; color:{_TEXT};
    padding:6px 10px; font-size:12px;
    min-width:120px;
}}
QDateEdit:focus {{ border:1px solid {_CYAN}; }}
QDateEdit:hover {{ border:1px solid {_CYAN}; }}
QDateEdit::drop-down {{
    border:none; background:transparent; width:28px;
    subcontrol-origin:border; subcontrol-position:right center;
}}
QDateEdit::down-arrow {{
    image:none; border-left:5px solid transparent;
    border-right:5px solid transparent;
    border-top:6px solid {_CYAN}; width:0; height:0;
    margin-right:8px;
}}

/* ── Calendar popup ──────────────────────────────────── */
QCalendarWidget {{
    background:{_CARD}; border:1px solid {_CYAN}; border-radius:8px;
    padding:8px;
}}
QCalendarWidget QWidget {{ background:{_CARD}; color:{_TEXT}; }}
QCalendarWidget QAbstractItemView {{
    background:{_CARD}; color:{_TEXT};
    selection-background-color:{_CYAN};
    alternate-background-color:{_ACCENT};
}}
QCalendarWidget QToolButton {{
    color:{_CYAN}; background:{_ACCENT};
    font-size:12px; font-weight:bold;
    border-radius:4px;
}}
QCalendarWidget QToolButton:hover {{ background:{_CYAN}44; border-radius:4px; }}
QCalendarWidget #qt_calendar_navigationbar {{
    background:{_ACCENT}; border-bottom:1px solid {_CYAN};
    border-radius:8px 8px 0 0;
}}
QCalendarWidget QSpinBox {{
    background:{_ACCENT}; color:{_TEXT}; border:1px solid {_CYAN};
    border-radius:4px; padding:2px 6px;
}}
QCalendarWidget QMenu {{
    background-color: #000000; color:{_TEXT};
    border:1px solid {_CYAN};
}}
QCalendarWidget QMenu::item {{
    background-color: #000000;
    color:{_TEXT};
}}
QCalendarWidget QMenu::item:selected {{
    background-color: {_CYAN}55;
}}
QCalendarWidget QHeaderView::section {{
    background:{_ACCENT}; color:{_CYAN};
    border:1px solid {_BORDER};
}}

/* ── Buttons ────────────────────────────────────────── */
QPushButton {{
    background:{_ACCENT}; color:{_CYAN};
    border:1px solid {_CYAN}55; border-radius:6px;
    padding:7px 16px; font-size:12px; font-weight:500;
}}
QPushButton:hover   {{ background:{_CYAN}1a; border:1px solid {_CYAN}; }}
QPushButton:pressed {{ background:{_CYAN}35; }}

QPushButton#primary {{
    background:qlineargradient(x1:0,y1:0,x2:1,y2:0,
        stop:0 {_CYAN}, stop:1 {_PINK});
    color:white; border:none; border-radius:8px;
    font-size:15px; font-weight:bold; padding:14px 0;
}}
QPushButton#primary:hover {{
    background:qlineargradient(x1:0,y1:0,x2:1,y2:0,
        stop:0 #20d8ff, stop:1 #f040e0);
}}
QPushButton#primary:pressed {{
    background:qlineargradient(x1:0,y1:0,x2:1,y2:0,
        stop:0 #00a8d0, stop:1 #b010b0);
}}

/* ── Table ──────────────────────────────────────────── */
QTableWidget {{
    background:{_ACCENT}; alternate-background-color:{_CARD};
    border:1px solid {_BORDER}; border-radius:6px;
    gridline-color:{_BORDER}; color:{_TEXT}; font-size:11px;
    selection-background-color:{_CYAN}30;
}}
QTableWidget::item:selected {{ background:{_CYAN}30; color:{_TEXT}; }}
QHeaderView::section {{
    background:#0d0d26; color:{_CYAN};
    border:none; border-bottom:1px solid {_CYAN}44;
    padding:6px 8px; font-weight:bold; font-size:11px;
}}
QScrollBar:vertical {{
    background:{_CARD}; width:7px; border-radius:3px;
}}
QScrollBar::handle:vertical {{
    background:{_CYAN}55; border-radius:3px; min-height:24px;
}}
QScrollBar::handle:vertical:hover {{ background:{_CYAN}; }}
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{ height:0; }}

/* ── Radio ──────────────────────────────────────────── */
QRadioButton {{ color:{_TEXT}; font-size:12px; spacing:8px; }}
QRadioButton::indicator {{
    width:15px; height:15px; border-radius:8px;
    border:2px solid {_CYAN}66; background:{_ACCENT};
}}
QRadioButton::indicator:checked {{
    border:2px solid {_CYAN};
    background:qradialgradient(cx:.5,cy:.5,radius:.45,fx:.5,fy:.5,
        stop:0 {_CYAN}, stop:.55 {_CYAN}55, stop:1 {_ACCENT});
}}

/* ── Title bar / Footer ─────────────────────────────── */
QFrame#titlebar {{
    background:qlineargradient(x1:0,y1:0,x2:1,y2:0,
        stop:0 #08081a, stop:.5 #0c0c22, stop:1 #08081a);
    border-bottom:1px solid {_CYAN}33;
}}
QFrame#footer {{
    background:{_CARD};
    border-top:1px solid {_BORDER};
}}

/* ── Total panel ────────────────────────────────────── */
QFrame#total_panel {{
    background:qlineargradient(x1:0,y1:0,x2:1,y2:0,
        stop:0 #001a0e, stop:1 #00101a);
    border:1px solid #003322; border-radius:6px;
}}

/* ── Settings form ──────────────────────────────────── */
QFormLayout QLabel {{ color:{_DIM}; }}
"""


class _Worker(QThread):
    done = pyqtSignal(dict)
    def __init__(self, fn, *a, **kw):
        super().__init__(); self._fn = fn; self._a = a; self._kw = kw
    def run(self): self.done.emit(self._fn(*self._a, **self._kw))


class _Card(QFrame):
    """Dark card with gradient header."""
    def __init__(self, icon: str, title: str, parent=None):
        super().__init__(parent)
        self.setObjectName("card")
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)
        hdr = QLabel(f"{icon}  {title}")
        hdr.setObjectName("card_hdr")
        root.addWidget(hdr)
        self._body = QWidget()
        self._body.setStyleSheet("background:transparent;")
        self._bl = QVBoxLayout(self._body)
        self._bl.setContentsMargins(14, 10, 14, 14)
        self._bl.setSpacing(8)
        root.addWidget(self._body)

    def bl(self): return self._bl


class App(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Happy Smart Light — Tạo Hợp Đồng Mua Bán")
        self.resize(1020, 880)
        self.setMinimumSize(840, 680)
        self.inv = None
        self._mst_cache: dict = {}

        logo_path = Path(__file__).parent / "logo.png"
        if logo_path.exists():
            self.setWindowIcon(QIcon(str(logo_path)))

        self._create_menubar()
        self._build()

    def _create_menubar(self):
        menubar = self.menuBar()
        
        about_menu = menubar.addMenu("ℹ️  Giới thiệu")
        
        about_action = about_menu.addAction("📋  Thông tin phiên bản")
        about_action.triggered.connect(self._show_about)
        
        about_menu.addSeparator()
        
        exit_action = about_menu.addAction("🚪  Thoát")
        exit_action.triggered.connect(self.close)

    def _show_about(self):
        from PyQt6.QtWidgets import QDialog, QVBoxLayout, QLabel, QPushButton
        from PyQt6.QtCore import Qt
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Thông tin")
        dialog.setFixedSize(420, 320)
        dialog.setStyleSheet(f"background:{_BG};")
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(16)
        layout.setContentsMargins(24, 24, 24, 24)
        
        logo_path = Path(__file__).parent / "logo.png"
        if logo_path.exists():
            pix_lbl = QLabel()
            pix = QPixmap(str(logo_path)).scaledToHeight(
                60, Qt.TransformationMode.SmoothTransformation)
            pix_lbl.setPixmap(pix)
            pix_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(pix_lbl)
        
        title = QLabel("Happy Smart Light — Tạo Hợp Đồng Mua Bán")
        title.setStyleSheet(f"color:{_CYAN}; font-size:16px; font-weight:bold; background:transparent;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)
        
        version = QLabel(f"Phiên bản: {VERSION}")
        version.setStyleSheet(f"color:{_TEXT}; font-size:13px; background:transparent;")
        version.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(version)
        
        info = QLabel(
            "Công cụ tạo hợp đồng mua bán từ hóa đơn điện tử XML/PDF\n"
            "Xuất file .docx chuẩn theo quy định Việt Nam"
        )
        info.setStyleSheet(f"color:{_DIM}; font-size:11px; background:transparent;")
        info.setAlignment(Qt.AlignmentFlag.AlignCenter)
        info.setWordWrap(True)
        layout.addWidget(info)
        
        layout.addSpacing(12)
        
        owner = QLabel(f"© Chủ sở hữu:\n{SELLER['name']}")
        owner.setStyleSheet(f"color:{_TEXT}; font-size:12px; background:transparent;")
        owner.setAlignment(Qt.AlignmentFlag.AlignCenter)
        owner.setWordWrap(True)
        layout.addWidget(owner)
        
        address = QLabel(f"📍 {SELLER['address']}")
        address.setStyleSheet(f"color:{_DIM}; font-size:10px; background:transparent;")
        address.setAlignment(Qt.AlignmentFlag.AlignCenter)
        address.setWordWrap(True)
        layout.addWidget(address)
        
        layout.addStretch()
        
        ok_btn = QPushButton("Đóng")
        ok_btn.setFixedHeight(36)
        ok_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        ok_btn.clicked.connect(dialog.close)
        layout.addWidget(ok_btn)
        
        dialog.exec()

    # ── Layout ────────────────────────────────────────────────
    def _build(self):
        root = QWidget(); self.setCentralWidget(root)
        vbox = QVBoxLayout(root)
        vbox.setContentsMargins(0, 0, 0, 0)
        vbox.setSpacing(0)

        vbox.addWidget(self._titlebar())

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        body = QWidget(); body.setStyleSheet(f"background:{_BG};")
        self._cl = QVBoxLayout(body)
        self._cl.setContentsMargins(16, 12, 16, 16)
        self._cl.setSpacing(12)
        scroll.setWidget(body)
        vbox.addWidget(scroll, 1)

        self._sec_invoice()
        self._sec_buyer()
        self._sec_goods()
        self._sec_contract()
        self._sec_bank()
        self._cl.addStretch()

        vbox.addWidget(self._footer())

    def _titlebar(self):
        bar = QFrame(); bar.setObjectName("titlebar"); bar.setFixedHeight(58)
        lay = QHBoxLayout(bar); lay.setContentsMargins(18, 0, 14, 0)

        logo_path = Path(__file__).parent / "logo.png"
        if logo_path.exists():
            pix_lbl = QLabel()
            pix = QPixmap(str(logo_path)).scaledToHeight(
                38, Qt.TransformationMode.SmoothTransformation)
            pix_lbl.setPixmap(pix)
            lay.addWidget(pix_lbl)
            lay.addSpacing(10)

        title = QLabel(f"⚡  Happy Smart Light — Tạo Hợp Đồng Mua Bán v{VERSION}")
        title.setStyleSheet(
            f"color:{_CYAN}; font-size:15px; font-weight:bold; background:transparent;")
        lay.addWidget(title)
        lay.addStretch()

        return bar

    def _footer(self):
        foot = QFrame(); foot.setObjectName("footer"); foot.setFixedHeight(72)
        lay = QHBoxLayout(foot); lay.setContentsMargins(20, 12, 20, 12)
        btn = QPushButton("📄   XUẤT HỢP ĐỒNG  (.docx)")
        btn.setObjectName("primary"); btn.setFixedHeight(46)
        btn.clicked.connect(self._export)
        self.btn_export = btn; lay.addWidget(btn)
        return foot

    # ── Helpers ───────────────────────────────────────────────
    def _row(self):
        w = QWidget(); w.setStyleSheet("background:transparent;")
        h = QHBoxLayout(w); h.setContentsMargins(0, 0, 0, 0); h.setSpacing(8)
        return w, h

    def _lbl(self, text, w=155):
        l = QLabel(text)
        l.setFixedWidth(w)
        l.setStyleSheet(f"color:{_DIM}; font-size:12px; background:transparent;")
        return l

    def _entry(self, ph="", fixed_w=0, min_w=0):
        e = QLineEdit(); e.setPlaceholderText(ph)
        if fixed_w: e.setFixedWidth(fixed_w)
        if min_w:   e.setMinimumWidth(min_w)
        return e

    def _date_edit(self, val: date = None):
        de = QDateEdit()
        de.setCalendarPopup(True)
        de.setDisplayFormat("dd/MM/yyyy")
        de.setDate(QDate(val.year, val.month, val.day) if val
                   else QDate.currentDate())
        de.setFixedWidth(120)
        return de

    def _badge(self, name="info") -> QLabel:
        l = QLabel(""); l.setObjectName(name); return l

    # ── Sections ──────────────────────────────────────────────
    def _sec_invoice(self):
        card = _Card("📄", "BƯỚC 1 — Chỉ nhận file hóa đơn nháp định dạng html (cấu trúc file mẫu :1C26TSL_0_3502535621.html)")
        self._cl.addWidget(card)
        rw, r = self._row()
        btn = QPushButton("📂  Chọn file")
        btn.setFixedSize(130, 34); btn.clicked.connect(self._pick_file)
        r.addWidget(btn)
        self.lbl_file = QLabel("Chưa chọn file…")
        self.lbl_file.setObjectName("dim"); r.addWidget(self.lbl_file, 1)
        card.bl().addWidget(rw)
        self.lbl_inv = QLabel(""); self.lbl_inv.setWordWrap(True)
        card.bl().addWidget(self.lbl_inv)

    def _sec_buyer(self):
        card = _Card("🏢", "BƯỚC 2 — Thông tin Bên Mua (BÊN B)")
        self._cl.addWidget(card)
        bl = card.bl()

        # MST
        mw, mr = self._row()
        mr.addWidget(self._lbl("Mã số thuế:"))
        self.e_mst = self._entry("Nhập MST…", fixed_w=190)
        self.e_mst.textChanged.connect(lambda: self._set_badge(self.lbl_mst, "", "info"))
        mr.addWidget(self.e_mst)
        btn_lk = QPushButton("🔍  Tra cứu")
        btn_lk.setFixedSize(110, 34); btn_lk.clicked.connect(self._lookup)
        mr.addWidget(btn_lk)
        self.lbl_mst = self._badge("info"); mr.addWidget(self.lbl_mst)
        mr.addStretch(); bl.addWidget(mw)

        # Name
        nw, nr = self._row()
        nr.addWidget(self._lbl("Tên công ty:"))
        self.e_buyer_name = self._entry("Tên đầy đủ…")
        nr.addWidget(self.e_buyer_name, 1); bl.addWidget(nw)

        # Address
        aw, ar = self._row()
        ar.addWidget(self._lbl("Địa chỉ:"))
        self.e_buyer_addr = self._entry("Địa chỉ…")
        ar.addWidget(self.e_buyer_addr, 1); bl.addWidget(aw)

        # Rep + Title
        rw2, rr = self._row()
        rr.addWidget(self._lbl("Đại diện:"))
        self.e_buyer_rep = self._entry("Họ và tên…", min_w=200)
        rr.addWidget(self.e_buyer_rep)
        rr.addWidget(self._lbl("  Chức vụ:", w=80))
        self.e_buyer_title = self._entry(fixed_w=150)
        self.e_buyer_title.setText("Giám đốc")
        rr.addWidget(self.e_buyer_title); rr.addStretch(); bl.addWidget(rw2)

    def _sec_goods(self):
        card = _Card("📦", "BƯỚC 3 — Hàng hóa (từ hóa đơn)")
        self._cl.addWidget(card)
        bl = card.bl()

        cols   = ("STT","Tên hàng hóa","ĐVT","SL","Đơn giá","Trước thuế","TS%","Tiền thuế","Trị giá TT")
        widths = (38, 260, 50, 40, 100, 115, 44, 100, 115)

        self.tbl = QTableWidget(0, len(cols))
        self.tbl.setHorizontalHeaderLabels(cols)
        self.tbl.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.tbl.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.tbl.setAlternatingRowColors(True)
        self.tbl.verticalHeader().setVisible(False)
        self.tbl.setFixedHeight(165)
        for i, w in enumerate(widths):
            self.tbl.setColumnWidth(i, w)
        self.tbl.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.Stretch)
        bl.addWidget(self.tbl)

        # Totals panel
        tp = QFrame(); tp.setObjectName("total_panel")
        tl = QHBoxLayout(tp); tl.setContentsMargins(14, 10, 14, 10)
        self.lbl_total = QLabel("Tổng thanh toán: —")
        self.lbl_total.setObjectName("total_amount")
        self.lbl_words = QLabel("")
        self.lbl_words.setObjectName("total_words")
        tl.addWidget(self.lbl_total); tl.addSpacing(14)
        tl.addWidget(self.lbl_words, 1)
        bl.addWidget(tp)

    def _sec_contract(self):
        card = _Card("📅", "BƯỚC 4 — Thông tin hợp đồng")
        self._cl.addWidget(card)
        bl = card.bl()

        # Row 1: số HĐ + ngày ký
        r1w, r1 = self._row()
        r1.addWidget(self._lbl("Số hợp đồng:"))
        self.e_contract_no = self._entry("vd: 01", fixed_w=70)
        r1.addWidget(self.e_contract_no)
        self.lbl_suffix = QLabel(f"/{date.today().year}/HDMB")
        self.lbl_suffix.setObjectName("suffix"); r1.addWidget(self.lbl_suffix)
        r1.addSpacing(18)
        r1.addWidget(self._lbl("Ngày ký:", w=78))
        self.de_sign = self._date_edit()
        self.de_sign.dateChanged.connect(self._on_sign_date)
        r1.addWidget(self.de_sign); r1.addStretch(); bl.addWidget(r1w)

        # Row 2: Tỉ lệ trả trước
        r2w, r2 = self._row()
        r2.addWidget(self._lbl("Tỉ lệ trả trước:"))
        self.combo_pay_ratio = QComboBox()
        self.combo_pay_ratio.addItems(["50%", "70%", "100%"])
        
        # Windows combo box fix: Use explicit QListView to prevent transparent popup
        list_view = QListView()
        list_view.setStyleSheet("background-color: #000000;")
        self.combo_pay_ratio.setView(list_view)
        
        self.combo_pay_ratio.currentTextChanged.connect(self._update_pay_labels)
        r2.addWidget(self.combo_pay_ratio)
        r2.addSpacing(10)
        r2.addWidget(self._lbl("Ngày TT đợt 1:", w=100))
        self.de_pay1 = self._date_edit()
        r2.addWidget(self.de_pay1)
        self.lbl_pay1 = QLabel(""); self.lbl_pay1.setObjectName("dim")
        r2.addWidget(self.lbl_pay1); r2.addStretch(); bl.addWidget(r2w)

        # Row 3: đợt 2 (chỉ hiển thị nếu không trả hết 100%)
        r3w, r3 = self._row()
        self.lbl_pay2_header = self._lbl("Ngày TT đợt 2 (còn lại):", w=185)
        r3.addWidget(self.lbl_pay2_header)
        self.de_pay2 = self._date_edit()
        r3.addWidget(self.de_pay2)
        self.lbl_pay2 = QLabel(""); self.lbl_pay2.setObjectName("dim")
        r3.addWidget(self.lbl_pay2); r3.addStretch(); bl.addWidget(r3w)

        # Row 4: giao hàng
        r4w, r4 = self._row()
        r4.addWidget(self._lbl("Ngày giao hàng:"))
        self.de_del = self._date_edit()
        r4.addWidget(self.de_del); r4.addStretch(); bl.addWidget(r4w)

    def _sec_bank(self):
        card = _Card("🏦", "BƯỚC 5 — Tài khoản nhận tiền (Bên Bán)")
        self._cl.addWidget(card)
        bl = card.bl()

        self.bank_grp = QButtonGroup(self)
        for i, acc in enumerate(SELLER["accounts"]):
            rb = QRadioButton(
                f"  STK  {acc['number']}   —   {acc['holder']}   —   {acc['bank']}")
            if i == 0: rb.setChecked(True)
            self.bank_grp.addButton(rb, i)
            bl.addWidget(rb)

        rb_custom = QRadioButton("  Tài khoản khác:")
        self.bank_grp.addButton(rb_custom, -1)
        rb_custom.toggled.connect(self._toggle_custom)
        bl.addWidget(rb_custom)

        # Custom frame
        cf = QWidget(); cf.setStyleSheet(
            f"background:{_ACCENT}; border-radius:6px;")
        cfl = QVBoxLayout(cf); cfl.setContentsMargins(12, 8, 12, 8); cfl.setSpacing(6)

        c1w, c1 = self._row()
        c1.addWidget(self._lbl("Số TK:", w=80))
        self.e_cno = self._entry("Số tài khoản", fixed_w=170); self.e_cno.setEnabled(False)
        c1.addWidget(self.e_cno)
        c1.addWidget(self._lbl("  Chủ TK:", w=80))
        self.e_cholder = self._entry("Tên chủ TK", min_w=200); self.e_cholder.setEnabled(False)
        c1.addWidget(self.e_cholder); c1.addStretch(); cfl.addWidget(c1w)

        c2w, c2 = self._row()
        c2.addWidget(self._lbl("Ngân hàng:", w=80))
        self.e_cbank = self._entry("Tên ngân hàng", min_w=340); self.e_cbank.setEnabled(False)
        c2.addWidget(self.e_cbank); c2.addStretch(); cfl.addWidget(c2w)

        bl.addWidget(cf)
        self._custom_entries = [self.e_cno, self.e_cholder, self.e_cbank]

    def _on_sign_date(self, qd: QDate):
        self.lbl_suffix.setText(f"/{qd.year()}/HDMB")

    def _update_pay_labels(self, ratio_text: str):
        # Hiển thị/ẩn trường ngày đợt 2 dựa trên tỉ lệ
        is_full = ratio_text == "100%"
        self.de_pay2.setVisible(not is_full)
        self.lbl_pay2_header.setVisible(not is_full)
        self.lbl_pay2.setVisible(not is_full)
        
        # Cập nhật nhãn tỉ lệ và số tiền
        if not self.inv:
            return
            
        tp = self.inv.total_payment
        if ratio_text == "100%":
            i1 = tp
            i2 = 0
            self.lbl_pay1.setText(f"= {fmt(i1)} đồng (trả hết)")
            # Ensure pay2 label is empty when hidden
            self.lbl_pay2.setText("")
        else:
            pct1 = int(ratio_text.replace("%", ""))
            i1 = round(tp * pct1 / 100)
            i2 = tp - i1
            self.lbl_pay1.setText(f"= {fmt(i1)} đồng ({pct1}%)")
            self.lbl_pay2.setText(f"= {fmt(i2)} đồng ({100 - pct1}%)")

    def _toggle_custom(self, checked: bool):
        for e in self._custom_entries:
            e.setEnabled(checked)

    # ── Events ────────────────────────────────────────────────
    def _pick_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Chọn file hóa đơn điện tử", "",
            "Hóa đơn điện tử (*.html);;All (*)")
        if not path: return
        self.lbl_file.setText(Path(path).name)
        self._set_badge(self.lbl_inv, "⏳  Đang đọc hóa đơn…", "warn")
        QApplication.processEvents()
        try:
            ext = Path(path).suffix.lower()
            if ext == ".xml":   self.inv = parse_xml(path)
            elif ext == ".pdf": self.inv = parse_pdf(path)
            elif ext == ".html": self.inv = parse_html(path)
            else: raise ValueError("Chỉ hỗ trợ HTML, XML và PDF")
            self._fill_invoice()
            self._set_badge(self.lbl_inv,
                f"✅  Hóa đơn số {self.inv.no}, ký hiệu {self.inv.serial}, ngày {self.inv.inv_date}",
                "ok")
        except Exception as e:
            self.inv = None
            self._set_badge(self.lbl_inv, f"❌  Lỗi: {e}", "bad")
            # Reset labels when error occurs
            self.lbl_total.setText("Tổng thanh toán: —")
            self.lbl_words.setText("")
            self.lbl_pay1.setText("")
            self.lbl_pay2.setText("")
            self.tbl.setRowCount(0)

    def _fill_invoice(self):
        if not self.inv: return
        self.e_mst.setText(self.inv.buyer_tax)
        self.e_buyer_name.setText(self.inv.buyer_name)
        self.e_buyer_addr.setText(self.inv.buyer_address)

        self.tbl.setRowCount(0)
        aligns = [
            Qt.AlignmentFlag.AlignCenter,
            Qt.AlignmentFlag.AlignLeft  | Qt.AlignmentFlag.AlignVCenter,
            Qt.AlignmentFlag.AlignCenter,
            Qt.AlignmentFlag.AlignCenter,
            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter,
            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter,
            Qt.AlignmentFlag.AlignCenter,
            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter,
            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter,
        ]
        for it in self.inv.items:
            r = self.tbl.rowCount(); self.tbl.insertRow(r)
            vals = [str(it["stt"]), it["name"], it["unit"], str(it["qty"]),
                    fmt(it["unit_price"]), fmt(it["before_tax"]),
                    it["tax_rate"], fmt(it["tax_amount"]), fmt(it["total"])]
            for col, (v, a) in enumerate(zip(vals, aligns)):
                item = QTableWidgetItem(v); item.setTextAlignment(a)
                self.tbl.setItem(r, col, item)

        tp = self.inv.total_payment
        self.lbl_total.setText(f"Tổng thanh toán: {fmt(tp)} đồng")
        self.lbl_words.setText(f"({self.inv.total_words})")
        
        # Cập nhật số tiền theo tỉ lệ hiện tại
        self._update_pay_labels(self.combo_pay_ratio.currentText())

        if self.inv.inv_date:
            try:
                d = datetime.strptime(self.inv.inv_date, "%Y-%m-%d").date()
                qd = QDate(d.year, d.month, d.day)
                for de in (self.de_sign, self.de_pay1, self.de_pay2, self.de_del):
                    de.setDate(qd)
            except:
                pass

    def _lookup(self):
        mst = self.e_mst.text().strip()
        if not mst:
            self._msg("Thiếu MST", "Vui lòng nhập mã số thuế trước.", "warning")
            return
        if mst in self._mst_cache:
            self._apply_mst(self._mst_cache[mst]); return

        self._set_badge(self.lbl_mst, "🔄  Đang tra cứu…", "warn")
        self._worker = _Worker(lookup_mst, mst)
        self._worker.done.connect(
            lambda r: (self._mst_cache.update({mst: r}), self._apply_mst(r)))
        self._worker.start()

    def _apply_mst(self, r: dict):
        if "error" in r:
            self._set_badge(self.lbl_mst, f"❌  {r['error']}", "bad")
        else:
            if r.get("status") == "active":
                self._set_badge(self.lbl_mst, "✅  Đang hoạt động", "ok")
            else:
                self._set_badge(self.lbl_mst, "⚠️  Đã ngưng hoạt động", "bad")
            # Cập nhật mã số thuế từ API
            if r.get("tax_id"):
                self.e_mst.setText(r["tax_id"])
            if r.get("name") and not self.e_buyer_name.text():
                self.e_buyer_name.setText(r["name"])
            if r.get("address") and not self.e_buyer_addr.text():
                self.e_buyer_addr.setText(r["address"])

    def _toggle_custom(self, checked: bool):
        for e in self._custom_entries:
            e.setEnabled(checked)

    def _set_badge(self, lbl: QLabel, text: str, kind: str):
        styles = {
            "ok":   f"color:{_GREEN}; font-size:11px; font-weight:bold; background:transparent;",
            "bad":  f"color:{_PINK};  font-size:11px; font-weight:bold; background:transparent;",
            "info": f"color:{_CYAN};  font-size:11px; font-weight:bold; background:transparent;",
            "warn": f"color:{_WARN};  font-size:11px; background:transparent;",
        }
        lbl.setStyleSheet(styles.get(kind, styles["info"]))
        lbl.setText(text)

    def _msg(self, title, text, kind="info"):
        box = QMessageBox(self)
        box.setWindowTitle(title); box.setText(text)
        box.setStyleSheet(_QSS)
        icons = {
            "info":    QMessageBox.Icon.Information,
            "warning": QMessageBox.Icon.Warning,
            "error":   QMessageBox.Icon.Critical,
        }
        box.setIcon(icons.get(kind, QMessageBox.Icon.Information))
        box.exec()

    # ── Validate & Export ─────────────────────────────────────
    def _validate(self):
        if not self.inv:                       return "Chưa tải hóa đơn. Vui lòng chọn file XML hoặc PDF."
        if not self.e_buyer_name.text().strip(): return "Chưa điền tên công ty Bên B."
        if not self.e_buyer_addr.text().strip(): return "Chưa điền địa chỉ Bên B."
        if not self.e_buyer_rep.text().strip():  return "Chưa điền tên người đại diện Bên B."
        if not self.e_contract_no.text().strip():return "Chưa điền số hợp đồng."
        return None

    def _export(self):
        err = self._validate()
        if err: self._msg("Thiếu thông tin", err, "warning"); return

        sel_id = self.bank_grp.checkedId()
        if sel_id == -1:
            bank = {"number": self.e_cno.text().strip(),
                    "holder": self.e_cholder.text().strip(),
                    "bank":   self.e_cbank.text().strip()}
            if not all(bank.values()):
                self._msg("Thiếu thông tin",
                          "Vui lòng điền đầy đủ thông tin tài khoản tùy chỉnh.", "warning")
                return
        else:
            bank = SELLER["accounts"][sel_id]

        buyer_short = re.sub(r"[^\w ]", "", self.e_buyer_name.text())[:18].strip()
        default_name = f"HD_{self.e_contract_no.text()}_{buyer_short}.docx"
        out_path, _ = QFileDialog.getSaveFileName(
            self, "Lưu hợp đồng", default_name, "Word Document (*.docx)")
        if not out_path: return

        def _qd2date(de: QDateEdit) -> date:
            q = de.date(); return date(q.year(), q.month(), q.day())

        try:
            ratio_text = self.combo_pay_ratio.currentText()
            pay_mode = "100" if ratio_text == "100%" else ("70_30" if ratio_text == "70%" else "50_50")
            
            data = {
                "contract_no":     self.e_contract_no.text().strip(),
                "sign_date":       _qd2date(self.de_sign),
                "pay1_date":       _qd2date(self.de_pay1),
                "pay2_date":       _qd2date(self.de_pay2),
                "delivery_date":   _qd2date(self.de_del),
                "buyer": {
                    "name":           self.e_buyer_name.text().strip(),
                    "address":        self.e_buyer_addr.text().strip(),
                    "tax_code":       self.e_mst.text().strip(),
                    "representative": self.e_buyer_rep.text().strip(),
                    "title":          self.e_buyer_title.text().strip(),
                },
                "items":            self.inv.items,
                "total_before_tax": self.inv.total_before_tax,
                "total_tax":        self.inv.total_tax,
                "total_payment":    self.inv.total_payment,
                "total_words":      self.inv.total_words,
                "bank":             bank,
                "pay_mode":         pay_mode,
            }
            generate_docx(data, out_path)
            self._msg("✅  Thành công", f"Hợp đồng đã xuất:\n{out_path}")
            try:
                os.startfile(out_path)
            except AttributeError:
                import subprocess; subprocess.Popen(["open", out_path])
        except Exception as e:
            self._msg("❌  Lỗi xuất file", str(e), "error")


# ── Auto-install & entry point ─────────────────────────────
def _ensure_deps():
    missing = []
    if not HAS_DOCX:     missing.append("python-docx")
    if not HAS_REQUESTS: missing.append("requests")
    if not HAS_BS4:      missing.extend(["beautifulsoup4", "lxml"])
    if missing:
        import subprocess
        print(f"📦  Đang cài đặt: {', '.join(missing)} …")
        subprocess.check_call([sys.executable, "-m", "pip", "install"] + missing)
        print("✅  Cài xong. Khởi động lại…")
        os.execv(sys.executable, [sys.executable] + sys.argv)

if __name__ == "__main__":
    _ensure_deps()
    app = QApplication(sys.argv)
    app.setStyleSheet(_QSS)
    window = App()
    window.show()
    sys.exit(app.exec())
