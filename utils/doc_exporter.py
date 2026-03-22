import os
from datetime import datetime, timedelta
from pathlib import Path
import logging

try:
    import docx
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as lxml_etree
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from models.product import ImportOrder
from models.cost_config import CostBreakdown
from utils.paths import get_resource_path
from utils.logger import get_logger

log = get_logger(__name__)

SELLER_INFO = {
    "name":           "CÔNG TY TNHH THƯƠNG MẠI VÀ CÔNG NGHỆ HAPPY SMART LIGHT",
    "address":        "42 Hà Đức Trọng, Phường Bà Rịa, Thành phố Hồ Chí Minh",
    "tax_code":       "3502535621",
    "representative": "NGUYỄN DUY BẰNG",
    "title":          "Giám đốc",
    "phone":          "0784140494",
    "email":          "happysmartlight@outlook.com",
    "website":        "https://happysmartlight.com/"
}

def set_cell(cell, text, bold=False, italic=False, sz=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, bg=None, color=None):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(text)
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(sz)
    run.font.name  = "Times New Roman"
    if color:
        run.font.color.rgb = color
    if bg:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg)
        tcPr.append(shd)

def fmt(n: float) -> str:
    return f"{int(n):,}".replace(",", ".")

def setup_header_watermark(sec):
    """Insert logo.png as a watermark behind text, mimicking the hop_dong_tool feature."""
    logo_path = get_resource_path("logo.png")
    if not logo_path.exists():
        return

    header = sec.header
    header.is_linked_to_previous = False
    hdr_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hdr_para.clear()
    
    # ── Header: Phone ─────────────────────────────────────────
    # We add phone number to the right of the header
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_hdr = hdr_para.add_run(f"📞 Hotline/Zalo: {SELLER_INFO['phone']}")
    r_hdr.font.size = Pt(10)
    r_hdr.font.name = "Times New Roman"
    r_hdr.font.color.rgb = RGBColor(120, 120, 120)

    # ── Watermark Logo ────────────────────────────────────────
    run = hdr_para.add_run()
    run.add_picture(str(logo_path), width=Cm(12))

    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A  = "http://schemas.openxmlformats.org/drawingml/2006/main"
    inline_elem = hdr_para._p.find(".//{%s}inline" % WP)
    if inline_elem is not None:
        cx_val = inline_elem.find("{%s}extent" % WP).get("cx")
        cy_val = inline_elem.find("{%s}extent" % WP).get("cy")
        
        # Center horizontally and vertically (approximate for A4)
        pos_x  = str((7560000 - int(cx_val)) // 2)
        pos_y  = str((10692000 - int(cy_val)) // 2)

        graphic_elem = inline_elem.find(".//{%s}graphic" % A)
        graphic_xml  = lxml_etree.tostring(graphic_elem, encoding="unicode") if graphic_elem is not None else ""

        anchor_xml = (
            f'<wp:anchor distT="0" distB="0" distL="0" distR="0" '
            f'simplePos="0" relativeHeight="251658240" behindDoc="1" '
            f'locked="0" layoutInCell="1" allowOverlap="1" '
            f'xmlns:wp="{WP}">'
            f'<wp:simplePos x="0" y="0"/>'
            f'<wp:positionH relativeFrom="page"><wp:posOffset>{pos_x}</wp:posOffset></wp:positionH>'
            f'<wp:positionV relativeFrom="page"><wp:posOffset>{pos_y}</wp:posOffset></wp:positionV>'
            f'<wp:extent cx="{cx_val}" cy="{cy_val}"/>'
            f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            f'<wp:wrapNone/>'
            f'<wp:docPr id="100" name="WatermarkLogo"/>'
            f'<wp:cNvGraphicFramePr/>'
            f'{graphic_xml}'
            f'</wp:anchor>'
        )
        try:
             anchor_elem = lxml_etree.fromstring(anchor_xml)
             drawing = inline_elem.getparent()
             drawing.remove(inline_elem)
             drawing.append(anchor_elem)
        except Exception as e:
             log.error(f"Failed to modify watermark anchor: {e}")

    # Set image transparency (alpha)
    a_blip = hdr_para._p.find(".//{%s}blip" % A)
    if a_blip is not None:
        alpha_mod = lxml_etree.SubElement(a_blip, "{%s}alphaModFix" % A)
        alpha_mod.set("amt", "15000") # 15% opacity

def export_quotation(order: ImportOrder, breakdown: CostBreakdown, out_path: Path, customer_name: str = "Quý Khách Hàng"):
    """
    Export a beautiful Word Document quotation.
    """
    if not HAS_DOCX:
        raise RuntimeError("Please install python-docx to generate Word files.")

    doc = Document()
    sec = doc.sections[0]
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # Thêm Header và Watermark mờ ở nền
    setup_header_watermark(sec)

    # Thêm Footer (Website bên trái, Số trang bên phải)
    footer = sec.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""
        
    ftr_table = footer.add_table(rows=1, cols=2, width=sec.page_width)
    ftr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = ftr_table.rows[0].cells
    for c in (c1, c2):
        tcPr = c._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for b in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'none')
            borders.append(border)
        tcPr.append(borders)
    
    p_left = c1.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_web = p_left.add_run(f"🌐 Website: {SELLER_INFO['website']}")
    r_web.font.size = Pt(10)
    r_web.font.name = "Times New Roman"
    r_web.font.color.rgb = RGBColor(120, 120, 120)
    
    p_right = c2.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_pg = p_right.add_run("Trang ")
    r_pg.font.size = Pt(10)
    r_pg.font.name = "Times New Roman"
    r_pg.font.color.rgb = RGBColor(120, 120, 120)
    
    r_pg_num = p_right.add_run()
    r_pg_num.font.size = Pt(10)
    r_pg_num.font.name = "Times New Roman"
    r_pg_num.font.color.rgb = RGBColor(120, 120, 120)

    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    r_pg_num._r.append(fld1)
    
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = "PAGE"
    r_pg_num._r.append(instr)
    
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    r_pg_num._r.append(fld2)
    
    r_pg_total = p_right.add_run(" / ")
    r_pg_total.font.size = Pt(10)
    r_pg_total.font.name = "Times New Roman"
    r_pg_total.font.color.rgb = RGBColor(120, 120, 120)
    
    r_pg_max = p_right.add_run()
    r_pg_max.font.size = Pt(10)
    r_pg_max.font.name = "Times New Roman"
    r_pg_max.font.color.rgb = RGBColor(120, 120, 120)

    fld3 = OxmlElement('w:fldChar')
    fld3.set(qn('w:fldCharType'), 'begin')
    r_pg_max._r.append(fld3)
    
    instr2 = OxmlElement('w:instrText')
    instr2.set(qn('xml:space'), 'preserve')
    instr2.text = "NUMPAGES"
    r_pg_max._r.append(instr2)
    
    fld4 = OxmlElement('w:fldChar')
    fld4.set(qn('w:fldCharType'), 'end')
    r_pg_max._r.append(fld4)

    FONT = "Times New Roman"

    def add_para(text="", bold=False, italic=False, sz=12,
                 align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=6):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_after  = Pt(space_after)
        r = para.add_run(text)
        r.bold      = bold
        r.italic    = italic
        r.font.size = Pt(sz)
        r.font.name = FONT
        if color:
             r.font.color.rgb = color
        return para

    # Header section
    add_para(SELLER_INFO["name"], bold=True, sz=14, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_para(SELLER_INFO["address"], italic=True, space_after=2)
    add_para(f"MST: {SELLER_INFO['tax_code']}", space_after=2)
    
    now = datetime.now()
    tp_str = str(int(breakdown.selling_price_vnd))
    last_4_amount = tp_str[-4:] if len(tp_str) >= 4 else tp_str.zfill(4)
    quotation_id = f"{now.strftime('%Y%m%d%H%M')}-{last_4_amount}"
    add_para(f"Mã báo giá: {quotation_id}", space_after=20)

    add_para("BẢNG BÁO GIÁ SẢN PHẨM", bold=True, sz=18, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(22, 54, 92), space_after=4)
    add_para(f"Ngày {datetime.now().day:02d} tháng {datetime.now().month:02d} năm {datetime.now().year}", 
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_para(f"Kính gửi: {customer_name}", bold=True, sz=13, space_after=6)
    add_para("Chúng tôi, HAPPY SMART LIGHT chân thành cảm ơn Quý khách hàng đã quan tâm đến sản phẩm của chúng tôi.\nChúng tôi xin gửi đến Quý khách bảng báo giá chi tiết như sau:", space_after=12)

    # Data Table
    headers = ["STT", "Tên sản phẩm", "ĐVT", "Số lượng", "Đơn giá (VND)", "Thành tiền (VND)"]
    col_widths = [Cm(1.2), Cm(7.0), Cm(1.5), Cm(2.0), Cm(3.0), Cm(3.5)]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, title in enumerate(headers):
        set_cell(hdr_row.cells[i], title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="E0E0E0", color=RGBColor(0, 0, 0))

    try:
        # Prevent accessing missing breakdowns if sizes don't match
        line_bds = breakdown.line_breakdowns
    except AttributeError:
        line_bds = []

    for i, line in enumerate(order.lines, 1):
        row = table.add_row()
        
        # Get individual selling price from breakdown item if available
        if i - 1 < len(line_bds):
            unit_price_vnd = line_bds[i-1].selling_price_vnd
            total_price_vnd = line_bds[i-1].total_selling_price_vnd
        else:
            # Fallback if breakdown array length doesn't match
            unit_price_vnd = 0
            total_price_vnd = 0
            
        vals = [
            str(i),
            line.product.name,
            "Cái", # Default unit
            fmt(line.product.qty),
            fmt(unit_price_vnd),
            fmt(total_price_vnd)
        ]
        
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT
        ]
        
        for idx, (v, a) in enumerate(zip(vals, aligns)):
            set_cell(row.cells[idx], v, align=a)
            
    # Add Total Row
    total_row = table.add_row()
    total_row.cells[0].merge(total_row.cells[3])
    set_cell(total_row.cells[0], "TỔNG CỘNG ĐỀ XUẤT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, bg="F2F2F2")
    set_cell(total_row.cells[4], "")
    set_cell(total_row.cells[5], fmt(breakdown.selling_price_vnd), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, bg="F2F2F2", color=RGBColor(192, 0, 0))

    doc.add_paragraph()

    # Text number to words (simple stub placeholder or just general notes)
    # We will just write standard terms
    add_para("Điều khoản áp dụng:", bold=True, sz=12, space_after=4)
    add_para("- Báo giá bao gồm chi phí giao hàng theo thỏa thuận.", space_after=2)
    add_para("- Phương thức thanh toán: Chuyển khoản hoặc tiền mặt.", space_after=2)
    
    valid_until = datetime.now() + timedelta(days=15)
    valid_str = f"{valid_until.day:02d}/{valid_until.month:02d}/{valid_until.year}"
    add_para(f"- Báo giá có hiệu lực trong vòng 15 ngày kể từ ngày báo (Hạn cuối: {valid_str}).", space_after=12)

    # Signature block
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = sig_table.rows[0].cells
    
    # Remove borders from signature table
    for cell in (c1, c2):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for b in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'none')
            borders.append(border)
        tcPr.append(borders)

    set_cell(c1, "XÁC NHẬN CỦA KHÁCH HÀNG", bold=True)
    set_cell(c2, f"ĐẠI DIỆN {SELLER_INFO['name'].upper()}", bold=True)
    
    p1 = c1.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("(Ký, ghi rõ họ tên)")

    p2 = c2.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("GIÁM ĐỐC\n\n\n\n\n")
    p2.add_run(SELLER_INFO['representative']).bold = True

    doc.add_paragraph()
    qr_p = doc.add_paragraph()
    qr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lbl = qr_p.add_run("Quét mã QR Zalo OA để được hỗ trợ nhanh nhất:\n")
    r_lbl.font.size = Pt(11)
    r_lbl.font.name = FONT
    r_lbl.italic = True

    qr_path = get_resource_path("qrcode_with_logo.png")
    if qr_path.exists():
        r_img = qr_p.add_run()
        r_img.add_picture(str(qr_path), width=Cm(3.0))

    try:
        doc.save(out_path)
        log.info(f"Successfully exported Word quotation to {out_path}")
    except Exception as e:
        log.error(f"Failed to save Word document: {e}")
        raise e

    return out_path
