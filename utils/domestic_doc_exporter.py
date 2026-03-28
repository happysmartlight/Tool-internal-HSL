"""
utils/domestic_doc_exporter.py
Exports a domestic pricing calculation to a Word quotation for customers.
Reuses helpers from utils.doc_exporter to avoid duplication.
Requires: python-docx
"""
from datetime import datetime
from pathlib import Path

from utils.doc_exporter import (
    SELLER_INFO, HAS_DOCX, set_cell, fmt,
    setup_header_watermark,
)
from models.domestic_product import DomesticBreakdown
from utils.logger import get_logger

log = get_logger(__name__)

if HAS_DOCX:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


def export_domestic_quotation(
    breakdown: DomesticBreakdown,
    out_path: Path,
    customer_name: str = "Quý Khách Hàng",
) -> Path:
    """
    Generate a customer-facing Word quotation (domestic products).
    Only shows selling prices — no cost or margin information.
    """
    if not HAS_DOCX:
        raise RuntimeError("Vui lòng cài python-docx để xuất file Word.")

    doc = Document()
    sec = doc.sections[0]
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    setup_header_watermark(sec)
    _setup_footer(sec)

    FONT = "Times New Roman"

    def add_para(text="", bold=False, italic=False, sz=12,
                 align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=6):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_after = Pt(space_after)
        r = para.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(sz)
        r.font.name = FONT
        if color:
            r.font.color.rgb = color
        return para

    # ── Company header ────────────────────────────────────────
    add_para(SELLER_INFO["name"], bold=True, sz=14, space_after=2)
    add_para(SELLER_INFO["address"], italic=True, space_after=2)
    add_para(f"MST: {SELLER_INFO['tax_code']}  |  ĐT: {SELLER_INFO['phone']}  |  Email: {SELLER_INFO['email']}", space_after=2)

    # ── Quotation ID ──────────────────────────────────────────
    now = datetime.now()
    tp_str = str(int(breakdown.total_revenue_with_vat))
    last_4 = tp_str[-4:] if len(tp_str) >= 4 else tp_str.zfill(4)
    quotation_id = f"ND-{now.strftime('%Y%m%d%H%M')}-{last_4}"
    add_para(f"Mã báo giá: {quotation_id}", space_after=18)

    # ── Title ─────────────────────────────────────────────────
    add_para("BẢNG BÁO GIÁ SẢN PHẨM NỘI ĐỊA", bold=True, sz=18,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(22, 54, 92), space_after=4)
    add_para(
        f"Ngày {now.day:02d} tháng {now.month:02d} năm {now.year}",
        italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
    )

    add_para(f"Kính gửi: {customer_name}", bold=True, sz=13, space_after=6)
    add_para(
        "Chúng tôi xin trân trọng gửi đến Quý khách hàng bảng báo giá sản phẩm như sau:",
        space_after=12,
    )

    # ── Product table ─────────────────────────────────────────
    headers = ["STT", "Tên sản phẩm", "ĐVT", "Số lượng",
               "Đơn giá (VND)", "Thành tiền (VND)"]
    col_widths = [Cm(1.2), Cm(7.0), Cm(1.5), Cm(2.0), Cm(3.2), Cm(3.3)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, title in enumerate(headers):
        set_cell(hdr_row.cells[i], title, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER,
                 bg="1E3A5F", color=RGBColor(255, 255, 255))
        hdr_row.cells[i].width = col_widths[i]

    for idx, line in enumerate(breakdown.lines, 1):
        row = table.add_row()
        cells = row.cells
        # Apply alternating row color
        row_bg = "F0F4FA" if idx % 2 == 0 else None

        set_cell(cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER, bg=row_bg)
        set_cell(cells[1], line.product_name, align=WD_ALIGN_PARAGRAPH.LEFT, bg=row_bg)
        set_cell(cells[2], line.unit, align=WD_ALIGN_PARAGRAPH.CENTER, bg=row_bg)
        set_cell(cells[3], fmt(line.qty) if line.qty != int(line.qty) else str(int(line.qty)),
                 align=WD_ALIGN_PARAGRAPH.CENTER, bg=row_bg)
        set_cell(cells[4], fmt(line.unit_sell_with_vat_vnd),
                 align=WD_ALIGN_PARAGRAPH.RIGHT, bg=row_bg)
        set_cell(cells[5], fmt(line.total_sell_with_vat_vnd),
                 align=WD_ALIGN_PARAGRAPH.RIGHT, bg=row_bg)
        for i, w in enumerate(col_widths):
            cells[i].width = w

    # ── Total row ─────────────────────────────────────────────
    total_row = table.add_row()
    total_row.cells[0].merge(total_row.cells[4])
    set_cell(total_row.cells[0], "TỔNG CỘNG (đã bao gồm VAT)",
             bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
             bg="1E3A5F", color=RGBColor(255, 255, 255), sz=12)
    set_cell(total_row.cells[5], fmt(breakdown.total_revenue_with_vat),
             bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
             bg="1E3A5F", color=RGBColor(255, 255, 255), sz=12)

    # ── Terms ─────────────────────────────────────────────────
    add_para("", space_after=8)
    add_para("ĐIỀU KHOẢN & ĐIỀU KIỆN:", bold=True, sz=12, space_after=4)
    terms = [
        "• Báo giá có hiệu lực trong vòng 15 ngày kể từ ngày lập.",
        "• Giá trên đã bao gồm thuế VAT 10%.",
        "• Thanh toán: Chuyển khoản ngân hàng hoặc tiền mặt.",
        "• Thời gian giao hàng: Thoả thuận.",
    ]
    for term in terms:
        add_para(term, sz=11, space_after=3)

    # ── Signature block ───────────────────────────────────────
    add_para("", space_after=24)
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(sig_table)

    c_customer, c_seller = sig_table.rows[0].cells

    def sig_col(cell, title, name, role):
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(title)
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.name = FONT

        cell.add_paragraph("")  # spacer
        cell.add_paragraph("")

        p_name = cell.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_name = p_name.add_run(name)
        r_name.bold = True
        r_name.font.size = Pt(11)
        r_name.font.name = FONT

        p_role = cell.add_paragraph()
        p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_role = p_role.add_run(role)
        r_role.font.size = Pt(10)
        r_role.font.name = FONT

    sig_col(c_customer, "ĐẠI DIỆN KHÁCH HÀNG", "(Ký và ghi rõ họ tên)", "")
    sig_col(c_seller,
            f"ĐẠI DIỆN HAPPY SMART LIGHT",
            SELLER_INFO["representative"],
            SELLER_INFO["title"])

    doc.save(out_path)
    log.info("Domestic Word quotation exported to %s", out_path)
    return out_path


def _setup_footer(sec):
    """Add website (left) and page number (right) to footer."""
    footer = sec.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""

    ftr_table = footer.add_table(rows=1, cols=2, width=sec.page_width)
    ftr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = ftr_table.rows[0].cells
    _remove_table_borders(ftr_table)

    p_left = c1.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_web = p_left.add_run(f"🌐 {SELLER_INFO['website']}")
    r_web.font.size = Pt(10)
    r_web.font.name = "Times New Roman"
    r_web.font.color.rgb = RGBColor(120, 120, 120)

    p_right = c2.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_pg = p_right.add_run("Trang ")
    r_pg.font.size = Pt(10)
    r_pg.font.name = "Times New Roman"
    r_pg.font.color.rgb = RGBColor(120, 120, 120)

    # PAGE field
    r_num = p_right.add_run()
    r_num.font.size = Pt(10)
    r_num.font.name = "Times New Roman"
    r_num.font.color.rgb = RGBColor(120, 120, 120)
    for tag, txt in [("begin", None), ("instrText", "PAGE"), ("end", None)]:
        if tag == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = txt
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        r_num._r.append(el)


def _remove_table_borders(table):
    """Remove all borders from a table's cells."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right"]:
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "none")
                borders.append(b)
            tcPr.append(borders)
